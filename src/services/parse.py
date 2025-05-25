"""Document parsing service for extracting information from resumes."""

import logging
import re
from datetime import datetime
from pathlib import Path
from typing import Dict, Any, List, Optional, Tuple
import PyPDF2
from docx import Document
import pytesseract
from pdf2image import convert_from_path
import spacy
from src.config import config, CACHE_DIR
from src.models.database import db, Candidate

logger = logging.getLogger(__name__)

# Configure Tesseract if path is provided
if config.tesseract_cmd:
    pytesseract.pytesseract.tesseract_cmd = config.tesseract_cmd


class ResumeParser:
    """Service for parsing resume documents."""
    
    def __init__(self):
        # Load spaCy model for NER
        try:
            self.nlp = spacy.load("en_core_web_sm")
        except OSError:
            logger.warning("spaCy model not found. Installing...")
            import subprocess
            subprocess.run(["python", "-m", "spacy", "download", "en_core_web_sm"])
            self.nlp = spacy.load("en_core_web_sm")
        
        # Compile regex patterns
        self.email_pattern = re.compile(r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b')
        self.phone_pattern = re.compile(r'[\+]?[(]?[0-9]{1,3}[)]?[-\s\.]?[(]?[0-9]{1,4}[)]?[-\s\.]?[0-9]{1,4}[-\s\.]?[0-9]{1,9}')
        
    def parse_all_pending(self) -> Tuple[int, int]:
        """
        Parse all pending resumes in the database.
        
        Returns:
            Tuple of (success_count, error_count)
        """
        success_count = 0
        error_count = 0
        
        with db.get_session() as session:
            # Get unparsed candidates
            candidates = session.query(Candidate).filter_by(is_parsed=False).all()
            
            for candidate in candidates:
                try:
                    # Get file path
                    file_path = CACHE_DIR / candidate.email_id / candidate.resume_filename
                    
                    if not file_path.exists():
                        logger.error(f"Resume file not found: {file_path}")
                        candidate.parse_error = "File not found"
                        error_count += 1
                        continue
                    
                    # Parse resume
                    parsed_data = self.parse_resume(file_path)
                    
                    # Update candidate with parsed data
                    self._update_candidate_with_parsed_data(session, candidate, parsed_data)
                    
                    success_count += 1
                    logger.info(f"Successfully parsed: {candidate.resume_filename}")
                    
                except Exception as e:
                    logger.error(f"Error parsing resume {candidate.resume_filename}: {str(e)}")
                    candidate.parse_error = str(e)
                    error_count += 1
                    
                    db.log_processing(
                        session,
                        candidate.email_id,
                        'parse',
                        'failed',
                        str(e)
                    )
            
            session.commit()
        
        logger.info(f"Parsing completed: {success_count} success, {error_count} errors")
        return success_count, error_count
    
    def parse_resume(self, file_path: Path) -> Dict[str, Any]:
        """Parse a single resume file."""
        file_ext = file_path.suffix.lower()
        
        # Extract text based on file type
        if file_ext == '.pdf':
            text = self._extract_pdf_text(file_path)
        elif file_ext in ['.doc', '.docx']:
            text = self._extract_docx_text(file_path)
        else:
            raise ValueError(f"Unsupported file type: {file_ext}")
        
        # Extract information from text
        parsed_data = {
            'text': text,
            'name': self._extract_name(text),
            'email': self._extract_email(text),
            'phone': self._extract_phone(text),
            'skills': self._extract_skills(text),
            'education': self._extract_education(text),
            'experience': self._extract_experience(text)
        }
        
        return parsed_data
    
    def _extract_pdf_text(self, file_path: Path) -> str:
        """Extract text from PDF file."""
        text = ""
        
        try:
            # Try regular text extraction first
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n"
            
            # If text is too short, might be scanned - use OCR
            if len(text.strip()) < 100:
                logger.info(f"PDF appears to be scanned, using OCR: {file_path.name}")
                text = self._extract_pdf_text_ocr(file_path)
                
        except Exception as e:
            logger.error(f"Error extracting PDF text: {str(e)}")
            # Fallback to OCR
            text = self._extract_pdf_text_ocr(file_path)
        
        return text
    
    def _extract_pdf_text_ocr(self, file_path: Path) -> str:
        """Extract text from PDF using OCR."""
        try:
            # Convert PDF to images
            images = convert_from_path(file_path)
            
            # Extract text from each page
            text = ""
            for i, image in enumerate(images):
                page_text = pytesseract.image_to_string(image)
                text += f"\n--- Page {i+1} ---\n{page_text}"
            
            return text
            
        except Exception as e:
            logger.error(f"OCR extraction failed: {str(e)}")
            return ""
    
    def _extract_docx_text(self, file_path: Path) -> str:
        """Extract text from DOCX file."""
        try:
            doc = Document(file_path)
            text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
            
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += "\n" + cell.text
            
            return text
            
        except Exception as e:
            logger.error(f"Error extracting DOCX text: {str(e)}")
            return ""
    
    def _extract_name(self, text: str) -> Optional[str]:
        """Extract candidate name from resume text."""
        # Use spaCy NER
        doc = self.nlp(text[:1000])  # Check first 1000 chars
        
        for ent in doc.ents:
            if ent.label_ == "PERSON":
                # Likely the first person mentioned is the candidate
                return ent.text
        
        # Fallback: Look for common patterns
        lines = text.split('\n')
        for i, line in enumerate(lines[:10]):  # Check first 10 lines
            line = line.strip()
            if line and len(line.split()) <= 4:  # Likely a name
                # Check if it's not an email or phone
                if not self.email_pattern.search(line) and not self.phone_pattern.search(line):
                    return line
        
        return None
    
    def _extract_email(self, text: str) -> Optional[str]:
        """Extract email address from resume text."""
        matches = self.email_pattern.findall(text)
        return matches[0] if matches else None
    
    def _extract_phone(self, text: str) -> Optional[str]:
        """Extract phone number from resume text."""
        matches = self.phone_pattern.findall(text)
        if matches:
            # Clean up the phone number
            phone = re.sub(r'[^\d+]', '', matches[0])
            return phone
        return None
    
    def _extract_skills(self, text: str) -> List[str]:
        """Extract skills from resume text."""
        skills = []
        
        # Common skill keywords
        skill_keywords = [
            'python', 'java', 'javascript', 'c++', 'c#', 'ruby', 'php', 'swift',
            'react', 'angular', 'vue', 'node.js', 'django', 'flask', 'spring',
            'docker', 'kubernetes', 'aws', 'azure', 'gcp', 'git', 'jenkins',
            'sql', 'nosql', 'mongodb', 'postgresql', 'mysql', 'redis',
            'machine learning', 'deep learning', 'data science', 'ai',
            'agile', 'scrum', 'devops', 'ci/cd', 'microservices'
        ]
        
        text_lower = text.lower()
        
        # Find skills section
        skills_section = ""
        skills_start = max(
            text_lower.find("skills"),
            text_lower.find("technical skills"),
            text_lower.find("core competencies")
        )
        
        if skills_start != -1:
            skills_section = text_lower[skills_start:skills_start + 1000]
        
        # Extract skills
        for skill in skill_keywords:
            if skill in text_lower:
                skills.append(skill)
        
        # Also look for skills in bullet points
        lines = text.split('\n')
        for line in lines:
            line = line.strip().lower()
            if line.startswith(('•', '-', '*', '▪', '◦')):
                for skill in skill_keywords:
                    if skill in line:
                        if skill not in skills:
                            skills.append(skill)
        
        return list(set(skills))  # Remove duplicates
    
    def _extract_education(self, text: str) -> List[Dict[str, str]]:
        """Extract education information from resume text."""
        education = []
        
        # Education keywords
        edu_keywords = ['education', 'academic', 'qualification', 'degree']
        degree_keywords = ['bachelor', 'master', 'phd', 'doctorate', 'mba', 'b.s.', 'm.s.', 'b.a.', 'm.a.']
        
        text_lower = text.lower()
        lines = text.split('\n')
        
        # Find education section
        edu_start = -1
        for keyword in edu_keywords:
            idx = text_lower.find(keyword)
            if idx != -1:
                edu_start = idx
                break
        
        if edu_start != -1:
            # Extract next few lines after education header
            edu_text = text[edu_start:edu_start + 1000]
            edu_lines = edu_text.split('\n')[1:10]  # Skip header, take next lines
            
            for line in edu_lines:
                line = line.strip()
                if line:
                    # Check for degree keywords
                    for degree in degree_keywords:
                        if degree in line.lower():
                            education.append({
                                'degree': line,
                                'text': line
                            })
                            break
        
        return education
    
    def _extract_experience(self, text: str) -> List[Dict[str, str]]:
        """Extract work experience from resume text."""
        experience = []
        
        # Experience keywords
        exp_keywords = ['experience', 'employment', 'work history', 'professional experience']
        
        text_lower = text.lower()
        
        # Find experience section
        exp_start = -1
        for keyword in exp_keywords:
            idx = text_lower.find(keyword)
            if idx != -1:
                exp_start = idx
                break
        
        if exp_start != -1:
            # Extract experience section
            exp_text = text[exp_start:exp_start + 2000]
            
            # Look for year patterns (e.g., 2020-2023, 2020 - Present)
            year_pattern = re.compile(r'(\d{4})\s*[-–]\s*(\d{4}|present|current)', re.IGNORECASE)
            matches = year_pattern.finditer(exp_text)
            
            for match in matches:
                # Get the line containing the year range
                start_pos = exp_text.rfind('\n', 0, match.start()) + 1
                end_pos = exp_text.find('\n', match.end())
                if end_pos == -1:
                    end_pos = len(exp_text)
                
                job_line = exp_text[start_pos:end_pos].strip()
                
                experience.append({
                    'period': match.group(0),
                    'title': job_line,
                    'years': self._calculate_years(match.group(1), match.group(2))
                })
        
        return experience
    
    def _calculate_years(self, start_year: str, end_year: str) -> float:
        """Calculate years of experience."""
        try:
            start = int(start_year)
            if end_year.lower() in ['present', 'current']:
                end = datetime.now().year
            else:
                end = int(end_year)
            return max(0, end - start)
        except:
            return 0
    
    def _update_candidate_with_parsed_data(self, session, candidate: Candidate, parsed_data: Dict[str, Any]):
        """Update candidate record with parsed data."""
        # Update basic info
        candidate.candidate_name = parsed_data.get('name')
        candidate.candidate_email = parsed_data.get('email')
        candidate.candidate_phone = parsed_data.get('phone')
        candidate.resume_text = parsed_data.get('text', '')
        
        # Update extracted data
        candidate.skills = parsed_data.get('skills', [])
        candidate.education = parsed_data.get('education', [])
        candidate.experience = parsed_data.get('experience', [])
        
        # Update flags
        candidate.is_parsed = True
        candidate.processed_at = datetime.utcnow()
        
        # Log success
        db.log_processing(
            session,
            candidate.email_id,
            'parse',
            'success',
            f'Extracted: {len(parsed_data.get("skills", []))} skills'
        )


# Global parser instance
resume_parser = ResumeParser() 